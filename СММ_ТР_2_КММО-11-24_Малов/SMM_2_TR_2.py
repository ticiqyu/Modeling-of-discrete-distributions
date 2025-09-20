import numpy as np
import matplotlib.pyplot as plt
import math
import random
import scipy.stats
from docx import Document
lambda1 = 0.55
lambda2 = 1.87
lambda3 = 0.32
q1 = 0.34
q2 = 0.43
q3 = 0.23

def f(x):
    return q1*lambda1*math.exp(-lambda1*x)+q2*lambda2*math.exp(-lambda2*x)+q3*lambda3*math.exp(-lambda3*x)
def F(x):
    return 1 - q1*math.exp(-lambda1*x) - q2*math.exp(-lambda2*x) - q3*math.exp(-lambda3*x)
def G_1(y):
    return math.log(1-y)/(-lambda1)
def G_2(y):
    return math.log(1 - y) / (-lambda2)
def G_3(y):
    return math.log(1-y)/(-lambda3)


def superpos():
    ksi= []
    for _ in range(200):
        alpha_1 = random.random()
        alpha_2 = random.random()
        if alpha_1<q1:
            ksi.append(round(G_1(alpha_2),5))
        elif alpha_1<q1+q2:
            ksi.append(round(G_2(alpha_2),5))
        else:
            ksi.append(round(G_3(alpha_2),5))
    return ksi

random.seed(1)
ksi_sup = superpos()
print(ksi_sup)

doc = Document()
doc.add_heading('Таблица случайных чисел (20x10)', level=1)
table = doc.add_table(rows=20, cols=10)
for i in range(20):
    for j in range(10):
        index = i * 10 + j
        table.cell(i, j).text = str(ksi_sup[index])
doc.save('случайные_числа.docx')
print("Документ создан и сохранен как 'случайные_числа.docx'")

ksi_sup.sort()
print(ksi_sup)

doc = Document()
doc.add_heading('Таблица случайных чисел (20x10)', level=1)
table = doc.add_table(rows=20, cols=10)
for i in range(20):
    for j in range(10):
        index = i * 10 + j
        table.cell(i, j).text = str(ksi_sup[index])
doc.save('случайные_числа.docx')
print("Документ создан и сохранен как 'случайные_числа.docx'")


import numpy as np
import math

def my_exp(x, lamb=0.81):
    return 1 - q1*math.exp(-lambda1*x) - q2*math.exp(-lambda2*x) - q3*math.exp(-lambda3*x)

def compute_intervals(data, m):
    a_0 = 0
    a_m = max(data)
    interval = (a_m - a_0) / m
    a_i = [round(a_0 + i * interval, 5) for i in range(1, m + 1)]
    return a_i

def compute_frequencies(data, a_i):
    n_i = np.zeros(len(a_i))
    for x in data:
        i = 0
        while i < len(a_i) - 1 and x > a_i[i]:
            i += 1
        n_i[i] += 1
    w_i = [round(n_i[i] / len(data), 5) for i in range(len(n_i))]
    return n_i, w_i

def compute_probabilities(a_i, func):
    prob = []
    prev = 0
    for x in a_i:
        current = func(x)
        prob.append(round(current - prev, 5))
        prev = current
    return prob

# Пример данных (замените на ваши ksi_exp и data_exp)

m = 8

print("KSI_SUP")
a_i_sup = compute_intervals(ksi_sup, m)
n_i_sup, w_i_sup = compute_frequencies(ksi_sup, a_i_sup)
print("Интервалы:", a_i_sup)
print("Частоты n_i:", n_i_sup)
print("Относительные частоты w_i:", w_i_sup, "\n")


prob = compute_probabilities(a_i_sup, my_exp)
print("prob_1 =", prob)
print("Сумма prob_1 =", sum(prob))




ksi_hyperexp_dif = [round(w_i_sup[i]-prob[i],7) for i in range(len(w_i_sup))]
print(ksi_hyperexp_dif)

hi_ksi = [200*ksi_hyperexp_dif[i]**2/prob[i] for i in range(len(ksi_hyperexp_dif))]
print(hi_ksi)
print(f"Хи-квадрат ДСП {sum(hi_ksi)}")
scipy.stats.chi2.ppf(0.95,m-1)

plt.bar([x for x in range(len(w_i_sup))],w_i_sup,color = 'Orange',label = 'Относительные частоты от метода ДСП')
plt.plot([x for x in range(len(prob))],prob,label = 'Теоретическое распределение плотностей')
plt.title("Гистограмма отн. частот ДСП")
plt.xlabel("i")
plt.ylabel("w_i_sup")
plt.grid()
plt.legend()
plt.show()

for a in a_i_sup:
    print(round(F(a),5))

